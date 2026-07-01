# -*- coding: utf-8 -*-
"""Génère le dossier technique Alim Express au format .docx avec la charte de l'app."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Palette Le Bakkal Oriental ----
GREEN     = RGBColor(0x3F,0x56,0x1F)
INK       = RGBColor(0x17,0x17,0x17)
MUTED     = RGBColor(0x6B,0x66,0x5D)
ALERT     = RGBColor(0xD5,0x2B,0x14)
WHITE     = RGBColor(0xFF,0xFF,0xFF)
GREEN_HEX = "3F561F"; SECTION_HEX="F0EBDD"; HILITE_HEX="F2C400"; ALERT_HEX="D52B14"
BG_HEX    = "FAF8F2"; BORDER_HEX="DDD8CC"; GREENSOFT_HEX="EAEDE3"
BODY_FONT = "Poppins"; HEAD_FONT = "Poppins"

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

def _shade(el, hex_color):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"),"clear")
    sh.set(qn("w:color"),"auto"); sh.set(qn("w:fill"),hex_color)
    el.append(sh)

def cell_bg(cell, hex_color):
    _shade(cell._tc.get_or_add_tcPr(), hex_color)

def para_shade(p, hex_color):
    _shade(p._p.get_or_add_pPr(), hex_color)

def set_borders(obj_pr, hex_color="DDD8CC", sz="8", space="6", sides=("top","left","bottom","right")):
    pbdr = OxmlElement("w:pBdr")
    for s in sides:
        b = OxmlElement(f"w:{s}")
        b.set(qn("w:val"),"single"); b.set(qn("w:sz"),sz)
        b.set(qn("w:space"),space); b.set(qn("w:color"),hex_color)
        pbdr.append(b)
    obj_pr.append(pbdr)

def run(p, text, *, size=10.5, bold=False, color=INK, font=BODY_FONT, italic=False):
    r = p.add_run(text); r.font.name=font; r.font.size=Pt(size)
    r.bold=bold; r.italic=italic; r.font.color.rgb=color
    return r

def heading(text, level=2):
    p = doc.add_paragraph(); p.paragraph_format.space_before=Pt(16 if level==2 else 10)
    p.paragraph_format.space_after=Pt(6)
    sizes={1:20,2:16,3:13,4:11}
    run(p, text, size=sizes[level], bold=True,
        color=(GREEN if level<=3 else INK), font=HEAD_FONT)
    if level==2:
        # underline bar via bottom border
        set_borders(p._p.get_or_add_pPr(), hex_color=HILITE_HEX, sz="18", space="4", sides=("bottom",))
    return p

def para(text=None, *, size=10.5, color=INK, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after=Pt(space_after)
    if text is not None:
        run(p, text, size=size, bold=bold, color=color, italic=italic)
    return p

def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after=Pt(3)
        if isinstance(it,(list,tuple)):
            run(p, it[0], bold=True); run(p, it[1])
        else:
            run(p, it)

def callout(parts):
    p = doc.add_paragraph(); pPr=p._p.get_or_add_pPr()
    para_shade(p, SECTION_HEX)
    set_borders(pPr, hex_color=GREEN_HEX, sz="24", space="8", sides=("left",))
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(10)
    p.paragraph_format.left_indent=Pt(10); p.paragraph_format.right_indent=Pt(6)
    for txt,bold in parts:
        run(p, txt, color=MUTED if not bold else GREEN, bold=bold, size=10)

def code_block(text):
    p = doc.add_paragraph(); para_shade(p, "23301A")
    p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(10)
    p.paragraph_format.left_indent=Pt(8); p.paragraph_format.right_indent=Pt(8)
    for i,line in enumerate(text.split("\n")):
        r=p.add_run(("" if i==0 else "\n")+line)
        r.font.name="Consolas"; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xED,0xE7,0xD6)

def placeholder(kind, title, desc, ref):
    accent = GREEN_HEX if kind=="diagram" else ALERT_HEX
    tag = "SCHÉMA À INSÉRER" if kind=="diagram" else "CAPTURE À INSÉRER"
    p = doc.add_paragraph(); pPr=p._p.get_or_add_pPr()
    para_shade(p, SECTION_HEX)
    set_borders(pPr, hex_color=accent, sz="12", space="10", sides=("top","left","bottom","right"))
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(10)
    p.paragraph_format.left_indent=Pt(8); p.paragraph_format.right_indent=Pt(8)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run(p, tag+"\n", bold=True, size=8.5, color=(GREEN if kind=="diagram" else ALERT), font=HEAD_FONT)
    run(p, title+"\n", bold=True, size=11, color=INK, font=HEAD_FONT)
    run(p, desc+"\n", size=9.5, color=MUTED)
    run(p, ref, size=9, color=GREEN, bold=True)

def table(headers, rows, *, kv=False, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.style="Table Grid"
    t.autofit=True
    hdr=t.rows[0].cells
    for i,h in enumerate(headers):
        cell_bg(hdr[i], GREEN_HEX)
        pp=hdr[i].paragraphs[0]; pp.paragraph_format.space_after=Pt(2)
        run(pp, h, bold=True, color=WHITE, size=9.5, font=HEAD_FONT)
    for ri,r in enumerate(rows):
        cells=t.add_row().cells
        for ci,val in enumerate(r):
            if ri%2==1: cell_bg(cells[ci], GREENSOFT_HEX)
            pp=cells[ci].paragraphs[0]; pp.paragraph_format.space_after=Pt(2)
            bold = (kv and ci==0)
            run(pp, val, size=9.5, bold=bold, color=(GREEN if bold else INK))
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

def crit(score, title):
    p=heading(f"{title}", level=2)
    # prepend score chip
    p.clear() if hasattr(p,"clear") else None
    # rebuild: chip + title
    for r in list(p.runs):
        r.text=""
    rc=p.add_run(f" {score} "); rc.font.name=HEAD_FONT; rc.bold=True; rc.font.size=Pt(12)
    rc.font.color.rgb=GREEN
    # shade the chip run isn't trivial; emphasize with highlight via WD highlight
    from docx.enum.text import WD_COLOR_INDEX
    rc.font.highlight_color=WD_COLOR_INDEX.YELLOW
    rt=p.add_run("  "+title); rt.font.name=HEAD_FONT; rt.bold=True; rt.font.size=Pt(16)
    rt.font.color.rgb=GREEN

# ============ COUVERTURE ============
def cover():
    # green block simulated via shaded full-width paragraphs
    sp=doc.add_paragraph(); sp.paragraph_format.space_after=Pt(0)
    for _ in range(2): doc.add_paragraph()
    p=doc.add_paragraph(); para_shade(p,GREEN_HEX)
    p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(2)
    p.paragraph_format.left_indent=Pt(8)
    run(p,"MBA DÉVELOPPEUR FULL STACK 2 · MYDIGITALSCHOOL · WORKSHOP CLIENT",
        bold=True,size=9,color=RGBColor(0xF2,0xC4,0x00),font=HEAD_FONT)
    p2=doc.add_paragraph(); para_shade(p2,GREEN_HEX)
    p2.paragraph_format.left_indent=Pt(8); p2.paragraph_format.space_after=Pt(2)
    run(p2,"Dossier technique",bold=True,size=40,color=WHITE,font=HEAD_FONT)
    p3=doc.add_paragraph(); para_shade(p3,GREEN_HEX)
    p3.paragraph_format.left_indent=Pt(8); p3.paragraph_format.space_after=Pt(6)
    run(p3,"Alim Express — plateforme de gestion et portail client pour un grossiste "
           "alimentaire halal-oriental, région PACA",size=13,color=RGBColor(0xED,0xE7,0xD6))
    p4=doc.add_paragraph(); para_shade(p4,GREEN_HEX)
    p4.paragraph_format.left_indent=Pt(8); p4.paragraph_format.space_after=Pt(14)
    run(p4,"  Enseigne · Le Bakkal Oriental  ",bold=True,size=10,color=GREEN,font=HEAD_FONT)
    from docx.enum.text import WD_COLOR_INDEX
    p4.runs[0].font.highlight_color=WD_COLOR_INDEX.YELLOW
    doc.add_paragraph()
    table(["Information","Valeur"],[
        ["Commanditaire","Rashed Younes, dirigeant d'Alim Express"],
        ["Encadrante pédagogique","Camille SIX"],
        ["Équipe","David · Anas · Yassin"],
        ["Période","3 mars → 22 juin 2026 (gel le 19 juin)"],
        ["Soutenance","1er juillet 2026"],
        ["Filière","Manager de projet web et digital"],
    ], kv=True)
    doc.add_page_break()

cover()

callout([("Objet de ce dossier. ",True),
    ("Conformément au guide du workshop client, ce document n'est pas un exercice de rédaction : "
     "il compile et donne à voir les livrables réellement produits (code source, outils de pilotage, "
     "documentation, recette) et en fournit les liens d'accès. Il est organisé selon les cinq critères "
     "de la grille « Réalisation du projet — Développeur full stack ».",False)])

# ============ IDENTITÉ ============
heading("Identité du projet",2)
table(["Élément","Valeur"],[
    ["Projet","Alim Express — back-office de gestion + portail client (B2C / B2B)"],
    ["Commanditaire","Rashed Younes, dirigeant d'Alim Express (grossiste halal-oriental, PACA)"],
    ["Enseigne (storefront)","Le Bakkal Oriental"],
    ["Formation","MBA Développeur Full Stack 2 — MyDigitalSchool"],
    ["Encadrante","Camille SIX"],
    ["Période / Gel de version","3 mars → 22 juin 2026 / 19 juin 2026"],
    ["Rendu / Soutenance","22 juin 2026 / 1er juillet 2026"],
], kv=True)

heading("Équipe et apports individuels",3)
table(["Membre","Rôle","Périmètre de réalisation"],[
    ["Yassin","Développeur back-end (partenaire technique du commanditaire)",
     "Cœur du back-office : produits, stock par lots & palettes (FEFO), commandes pro, devis, avoirs, tableau de bord ; explorations IA."],
    ["Anas","Développeur front-end",
     "Portail client intégral : parcours particuliers et professionnels, UX responsive mobile-first, PDF devis/facture côté client."],
    ["David","Coordination projet + développeur",
     "API reliant back-office et portail ; modules clients particuliers, facturation, livraisons ; pilotage, relation commanditaire, note de cadrage, recettage."],
])

# ============ INDEX LIVRABLES ============
heading("Index des livrables",2)
para("Liens d'accès pour le jury — à compléter avec les URL réelles avant le rendu.",
     italic=True, color=MUTED)
table(["Livrable","Type","Accès"],[
    ["Note de cadrage (v1.0 approuvée 03/03/2026)","PDF","docs/doc/Note-de-cadrage-Alim-Express-David.pdf"],
    ["Code source — back-office (API + admin)","GitHub","github.com/mineboxarabic/AlimExpressApp"],
    ["Code source — portail client (B2C + B2B)","GitHub","github.com/mineboxarabic/albakkaloriental_front"],
    ["Tableau de pilotage (backlog, sprints)","Trello","lien partagé en consultation"],
    ["Environnement de démonstration","App déployée","URL de démo"],
    ["Documentation de tests","Markdown","docs/doc/TESTING.md"],
    ["Procès-verbal de recettage","PDF signé","à joindre"],
    ["Guides utilisateur (back-office, portail)","PDF","à joindre"],
])

# ============ CRITÈRE 1 ============
crit("/ 5","Critère 1 — Structuration technique de l'application")
heading("1.1 Architecture : deux applications, une base partagée",3)
para("L'architecture repose sur deux applications web distinctes partageant une même base PostgreSQL. "
     "Le portail consomme l'API du back-office et n'accède pas directement à la base — ce qui isole la "
     "surface publique de la surface interne et simplifie le raisonnement sur la sécurité.")
placeholder("diagram","Schéma 1 — Architecture applicative",
    "Deux apps (portail client ↔ back-office via API REST) + base PostgreSQL partagée.",
    "→ Prompt de génération fourni séparément (Schéma 1)")

heading("1.2 Stack technique",3)
table(["Couche","Back-office (AlimExpressApp)","Portail (AlimExpressCatalog)"],[
    ["Framework","Next.js 15 (App Router), React 19","Next.js 15 (App Router), React 19"],
    ["Langage","TypeScript strict","TypeScript strict"],
    ["UI","Tailwind CSS + shadcn/ui","Tailwind CSS v4 + shadcn/ui + Lucide"],
    ["Données","PostgreSQL + Prisma (~980 lignes)","API REST /api/v1/*"],
    ["Authentification","Session serveur (admin) + JWT (clients)","JWT portable (jose), cookie HttpOnly"],
    ["Validation","Zod","Zod"],
    ["Documents","Génération PDF (devis, factures, bons)","Rendu PDF devis/facture"],
    ["Tests","Vitest + Playwright","Vitest + Testing Library + Playwright"],
])

heading("1.3 Organisation du code — portail client",3)
code_block("""app/
├── (public)/   → CGV, confidentialité, contact, mots de passe, vérification e-mail
├── (retail)/   → B2C : accueil, produits, panier, checkout, compte, suivi commande
└── (pro)/      → B2B : login pro, catalogue à niveaux tarifaires, devis, factures
actions/        → Server Actions validées par Zod (pro-quote.ts, pro-order.ts, …)
lib/            → api-client.ts, session.ts, catalog-pricing.ts (C/D/E/F), order-rules.ts
middleware.ts   → protection des routes /pro/* (vérification JWT)""")
para("Les groupes de routes isolent les parcours B2C / B2B / public ; les règles métier sont "
     "centralisées pour éviter la duplication (DRY).")

heading("1.4 Authentification — trois publics, trois mécanismes",3)
bullets([
    ["Administrateurs : ","session côté serveur (révocation immédiate)."],
    ["Particuliers & professionnels : ","JWT portable émis par le back-office, vérifié par le portail, en cookie HttpOnly."],
    ["Cloisonnement strict : ","le portail ne peut jamais atteindre une opération d'administration."],
])

heading("1.5 Cycle commercial & points techniques notables",3)
placeholder("diagram","Schéma 2 — Cycle commercial B2B",
    "Devis → acceptation horodatée → commande ferme → facture → avoir (remise en stock).",
    "→ Prompt de génération fourni séparément (Schéma 2)")
bullets([
    ["Tarification par niveaux (C/D/E/F) ","calculée dynamiquement, jamais exposée à un visiteur non authentifié."],
    ["Stock FEFO ","par lots/palettes, déduction automatique à la livraison."],
    ["Logistique : ","adresses géolocalisées, zonage, ordre de tournée, verrouillage J-1 à 21h."],
])

heading("1.6 Qualité logicielle",3)
bullets([
    "Couverture de tests > 60 % sur les modules critiques (commandes, stock, facturation, auth).",
    "Validation systématique aux frontières réseau (Zod), journal d'audit des opérations sensibles.",
    "Revue de code en binôme (Pull Requests), commits conventionnels.",
])

# ============ CRITÈRE 2 ============
crit("/ 3","Critère 2 — Contraintes d'hébergement anticipées")
para("L'environnement du commanditaire était minimal (tableurs + papier), sans contrainte "
     "d'intégration. Les choix ont privilégié la sobriété tarifaire et la fiabilité.")
table(["Poste","Service retenu","Tarification","Coût/an"],[
    ["Hébergement web + base de données","Railway (plan Hobby)","5 $/mois, crédits inclus","~60 €"],
    ["Stockage des PDF","AWS S3","à l'usage, faible volumétrie","~60 €"],
    ["E-mails transactionnels","Resend (gratuit)","3 000 e-mails/mois","0 €"],
    ["Notifications mobiles","Twilio (WhatsApp)","0,005 $/message","~40 €"],
    ["Géocodage des adresses","Référentiel public (BAN)","gratuit","0 €"],
    ["Total infrastructure v1","","","~160 €"],
])
heading("Anticipations intégrées au développement",4)
bullets([
    ["Variables d'environnement ","centralisées (.env.example) : BACKEND_URL, AUTH_SECRET (secret JWT partagé)."],
    ["Images produit ","servies depuis S3, domaine distant autorisé dans next.config.ts."],
    ["Mode dégradé ","prévu pour chaque dépendance tierce critique."],
    ["Montée en charge : ","indexation soignée, monitoring, tests de charge avant mise en production."],
    ["Cibles : ","page catalogue < 1 s (95e centile), erreurs serveur < 1 %, disponibilité > 99 %."],
])

# ============ CRITÈRE 3 ============
crit("/ 5","Critère 3 — Administration facilitée par un back-office")
para("L'administration est assurée par AlimExpressApp, qui centralise toute la gestion d'Alim Express.")
placeholder("capture","Capture 1 — Tableau de bord du back-office",
    "Vue d'ensemble de l'activité (commandes, stock, indicateurs).",
    "→ Source de capture fournie séparément (Capture 1)")
table(["Module","Fonctions principales"],[
    ["Produits","Catalogue, SKU/codes-barres, images S3, niveaux tarifaires C/D/E/F, visibilité RETAIL/WHOLESALE/BOTH."],
    ["Stock","Lots (péremption, FEFO), palettes, déduction auto à la livraison, traçabilité anti-survente."],
    ["Clients","Particuliers (auto-inscription) et pros (création manuelle + activation par lien)."],
    ["Cycle commercial","Devis, commande ferme, facture, avoir (remise en stock), documents PDF."],
    ["Livraisons","Géolocalisation, zonage, tournées, verrouillage J-1 21h, bons PDF."],
    ["Transverse","Tableau de bord, journal d'audit, notifications e-mail/WhatsApp, multilingue FR/AR."],
])
heading("Spécifications techniques et fonctionnelles écrites",4)
bullets([
    ["Modèle de données : ","schéma Prisma (~980 lignes), déployable sur environnement vierge, jeu de données rejouable."],
    ["Contrat d'API : ","points d'entrée /api/v1/* documentés, validations en entrée/sortie."],
    ["Spécifications fonctionnelles ","des modules (AGENTS.md) + documentation utilisateur (guide + formation 2 h)."],
])

# ============ CRITÈRE 4 ============
crit("/ 4","Critère 4 — Recettage formalisé et corrections")
para("Recettage à deux niveaux automatisés + recette manuelle avec le commanditaire, donnant lieu "
     "à un procès-verbal signé.")
bullets([
    ["Tests unitaires / composants ","— Vitest + Testing Library : logique lib/, Server Actions (back mocké), composants. Dossier __tests__/."],
    ["Tests end-to-end ","— Playwright : parcours complets contre un back mocké (e2e/mock-backend.ts). Dossier e2e/."],
    ["Objectif de couverture ","> 60 % sur les modules critiques."],
])
heading("Exemple de plan de test (E2E — parcours panier, e2e/cart.spec.ts)",4)
table(["#","Étape","Résultat attendu"],[
    ["1","Ouvrir le catalogue retail","Produits affichés"],
    ["2","Ajouter un produit au panier","Panier mis à jour"],
    ["3","Dépasser la quantité maximale","Limite par produit respectée"],
    ["4","Aller au paiement sous le minimum","Minimum de commande exigé"],
    ["5","Atteindre le seuil de franco","Frais de livraison offerts"],
    ["6","Valider la commande","Confirmation affichée"],
])
placeholder("capture","Capture 2 — Exécution de la suite de tests",
    "Sortie terminal Vitest + Playwright au vert (et/ou rapport de couverture).",
    "→ Source de capture fournie séparément (Capture 2)")
code_block("""npx vitest run --pool=threads      # tests unitaires et composants
npx playwright test                 # tests end-to-end
npm run lint                        # qualité du code (ESLint)""")
para("Corrections réalisées (traçables dans Git) : refonte de la navigation retail suite aux retours "
     "du commanditaire, ajustements responsive, intégration du branding Le Bakkal Oriental (logo, "
     "favicon). Le PV de recette formalise le périmètre testé, les anomalies, les corrections et le "
     "statut final de chaque parcours.")

# ============ CRITÈRE 5 ============
crit("/ 3","Critère 5 — Responsive et mobile-first")
para("Le portail est conçu en approche mobile-first : pensé et validé prioritairement sur smartphone, "
     "puis décliné sur tablette et ordinateur — reflet de l'usage réel (commande depuis le téléphone).")
placeholder("capture","Capture 3 — Portail responsive (mobile / desktop)",
    "Même page catalogue affichée côte à côte en vue mobile et vue desktop.",
    "→ Source de capture fournie séparément (Capture 3)")
heading("Charte graphique — Le Bakkal Oriental",4)
table(["Usage","Couleur","Hex"],[
    ["Vert principal (CTA, nav)","Olive foncé","#3F561F"],
    ["Fond principal","Blanc cassé chaud","#FAF8F2"],
    ["Fond de section","Beige crème","#F0EBDD"],
    ["Texte principal","Noir","#171717"],
    ["Texte secondaire","Gris chaud","#6B665D"],
    ["Bordures","Beige-gris clair","#DDD8CC"],
    ["Alerte / promo","Rouge-orangé","#D52B14"],
    ["Mise en avant","Jaune chaud","#F2C400"],
])
para("Typographie : Poppins (titres) et Satoshi (corps).")

# ============ GLOSSAIRE ============
heading("Annexe — Glossaire technique",2)
table(["Terme","Définition"],[
    ["B2C / B2B","Vente aux particuliers (prix standards) / aux professionnels (prix négociés par niveau)."],
    ["Devis (proforma)","Document chiffré pré-commande, soumis au client pour acceptation."],
    ["Commande ferme","Devis accepté et signé, déclenchant la préparation."],
    ["Avoir","Document de crédit avec réintégration du stock."],
    ["FEFO","First-Expired-First-Out : sortie du lot dont la péremption est la plus proche."],
    ["Niveau tarifaire","Catégorie de prix (C/D/E/F) attribuée à un client professionnel."],
    ["JWT","Jeton portable vérifiable entre le back-office et le portail."],
])
callout([("Pour le contexte marché, les enjeux stratégiques, le budget, le planning et les risques, "
          "se reporter à la note de cadrage (livrable joint).",False)])

out = r"c:\Users\anasl\Documents\AlimExpressCatalog\docs\doc\Dossier-technique-Alim-Express.docx"
doc.save(out)
print("OK ->", out)
