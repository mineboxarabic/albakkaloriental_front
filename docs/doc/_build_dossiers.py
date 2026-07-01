# -*- coding: utf-8 -*-
"""Génère deux dossiers techniques (.docx) — back-office & portail client — charte Le Bakkal Oriental."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN=RGBColor(0x3F,0x56,0x1F); INK=RGBColor(0x17,0x17,0x17); MUTED=RGBColor(0x6B,0x66,0x5D)
ALERT=RGBColor(0xD5,0x2B,0x14); WHITE=RGBColor(0xFF,0xFF,0xFF)
GREEN_HEX="3F561F"; SECTION_HEX="F0EBDD"; HILITE_HEX="F2C400"; ALERT_HEX="D52B14"
GREENSOFT_HEX="EAEDE3"; FONT="Poppins"

def _shade(el,hex_color):
    sh=OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:color"),"auto")
    sh.set(qn("w:fill"),hex_color); el.append(sh)
def _borders(pPr,hex_color="DDD8CC",sz="8",space="6",sides=("top","left","bottom","right")):
    pbdr=OxmlElement("w:pBdr")
    for s in sides:
        b=OxmlElement(f"w:{s}"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),sz)
        b.set(qn("w:space"),space); b.set(qn("w:color"),hex_color); pbdr.append(b)
    pPr.append(pbdr)

class Doc:
    def __init__(self):
        self.d=Document()
        n=self.d.styles["Normal"]; n.font.name=FONT; n.font.size=Pt(10.5)
        n.font.color.rgb=INK; n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.22
    def run(self,p,t,*,size=10.5,bold=False,color=INK,italic=False,mono=False):
        r=p.add_run(t); r.font.name=("Consolas" if mono else FONT); r.font.size=Pt(size)
        r.bold=bold; r.italic=italic; r.font.color.rgb=color; return r
    def h(self,t,level=2):
        p=self.d.add_paragraph(); p.paragraph_format.space_before=Pt(16 if level==2 else 10)
        p.paragraph_format.space_after=Pt(6); sizes={1:20,2:15,3:12.5,4:11}
        self.run(p,t,size=sizes[level],bold=True,color=(GREEN if level<=3 else INK))
        if level==2: _borders(p._p.get_or_add_pPr(),HILITE_HEX,"18","4",("bottom",))
        return p
    def p(self,t=None,*,size=10.5,color=INK,bold=False,italic=False,sa=6):
        p=self.d.add_paragraph(); p.paragraph_format.space_after=Pt(sa)
        if t is not None: self.run(p,t,size=size,bold=bold,color=color,italic=italic)
        return p
    def bullets(self,items):
        for it in items:
            p=self.d.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(2)
            if isinstance(it,(list,tuple)): self.run(p,it[0],bold=True); self.run(p,it[1])
            else: self.run(p,it)
    def callout(self,parts):
        p=self.d.add_paragraph(); pPr=p._p.get_or_add_pPr(); _shade(pPr,SECTION_HEX)
        _borders(pPr,GREEN_HEX,"24","8",("left",))
        p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(10)
        p.paragraph_format.left_indent=Pt(10); p.paragraph_format.right_indent=Pt(6)
        for txt,bold in parts: self.run(p,txt,color=(GREEN if bold else MUTED),bold=bold,size=10)
    def code(self,text):
        p=self.d.add_paragraph(); _shade(p._p.get_or_add_pPr(),"23301A")
        p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(10)
        p.paragraph_format.left_indent=Pt(8); p.paragraph_format.right_indent=Pt(8)
        for i,line in enumerate(text.split("\n")):
            r=p.add_run(("" if i==0 else "\n")+line); r.font.name="Consolas"; r.font.size=Pt(8.8)
            r.font.color.rgb=RGBColor(0xED,0xE7,0xD6)
    def ph(self,kind,title,desc,ref):
        accent=GREEN_HEX if kind=="diagram" else ALERT_HEX
        tag="SCHÉMA À INSÉRER" if kind=="diagram" else "CAPTURE À INSÉRER"
        col=GREEN if kind=="diagram" else ALERT
        p=self.d.add_paragraph(); pPr=p._p.get_or_add_pPr(); _shade(pPr,SECTION_HEX)
        _borders(pPr,accent,"12","10")
        p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(10)
        p.paragraph_format.left_indent=Pt(8); p.paragraph_format.right_indent=Pt(8)
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        self.run(p,tag+"\n",bold=True,size=8.5,color=col)
        self.run(p,title+"\n",bold=True,size=11,color=INK)
        self.run(p,desc+"\n",size=9.5,color=MUTED)
        self.run(p,ref,size=9,color=GREEN,bold=True)
    def table(self,headers,rows,*,kv=False):
        t=self.d.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
        t.style="Table Grid"
        for i,htext in enumerate(headers):
            c=t.rows[0].cells[i]; _shade(c._tc.get_or_add_tcPr(),GREEN_HEX)
            pp=c.paragraphs[0]; pp.paragraph_format.space_after=Pt(2)
            self.run(pp,htext,bold=True,color=WHITE,size=9.5)
        for ri,r in enumerate(rows):
            cells=t.add_row().cells
            for ci,val in enumerate(r):
                if ri%2==1: _shade(cells[ci]._tc.get_or_add_tcPr(),GREENSOFT_HEX)
                pp=cells[ci].paragraphs[0]; pp.paragraph_format.space_after=Pt(2)
                bold=(kv and ci==0); self.run(pp,val,size=9.3,bold=bold,color=(GREEN if bold else INK))
        self.d.add_paragraph().paragraph_format.space_after=Pt(2)
    def crit(self,score,title):
        p=self.d.add_paragraph(); p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(6)
        rc=p.add_run(f" {score} "); rc.font.name=FONT; rc.bold=True; rc.font.size=Pt(12)
        rc.font.color.rgb=GREEN; rc.font.highlight_color=WD_COLOR_INDEX.YELLOW
        rt=p.add_run("  "+title); rt.font.name=FONT; rt.bold=True; rt.font.size=Pt(15); rt.font.color.rgb=GREEN
        _borders(p._p.get_or_add_pPr(),HILITE_HEX,"18","4",("bottom",))
    def cover(self,kicker,title,subtitle,pill,meta):
        for _ in range(2): self.d.add_paragraph()
        p=self.d.add_paragraph(); _shade(p._p.get_or_add_pPr(),GREEN_HEX)
        p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(2); p.paragraph_format.left_indent=Pt(8)
        self.run(p,kicker,bold=True,size=9,color=RGBColor(0xF2,0xC4,0x00))
        p2=self.d.add_paragraph(); _shade(p2._p.get_or_add_pPr(),GREEN_HEX)
        p2.paragraph_format.left_indent=Pt(8); p2.paragraph_format.space_after=Pt(2)
        self.run(p2,title,bold=True,size=36,color=WHITE)
        p3=self.d.add_paragraph(); _shade(p3._p.get_or_add_pPr(),GREEN_HEX)
        p3.paragraph_format.left_indent=Pt(8); p3.paragraph_format.space_after=Pt(6)
        self.run(p3,subtitle,size=13,color=RGBColor(0xED,0xE7,0xD6))
        p4=self.d.add_paragraph(); _shade(p4._p.get_or_add_pPr(),GREEN_HEX)
        p4.paragraph_format.left_indent=Pt(8); p4.paragraph_format.space_after=Pt(14)
        r=p4.add_run("  "+pill+"  "); r.font.name=FONT; r.bold=True; r.font.size=Pt(10)
        r.font.color.rgb=GREEN; r.font.highlight_color=WD_COLOR_INDEX.YELLOW
        self.d.add_paragraph()
        self.table(["Information","Valeur"],meta,kv=True)
        self.d.add_page_break()
    def save(self,path): self.d.save(path); print("OK ->",path)

# ===================================================================================
# DOSSIER 1 — BACK-OFFICE (AlimExpressApp)
# ===================================================================================
def build_backoffice():
    b=Doc()
    b.cover("MBA DÉVELOPPEUR FULL STACK 2 · MYDIGITALSCHOOL · WORKSHOP CLIENT",
        "Dossier technique — Back-office",
        "Alim Express — application d'administration interne et API de la plateforme",
        "Enseigne · Le Bakkal Oriental",
        [["Application","AlimExpressApp (back-office + API REST)"],
         ["Commanditaire","Rashed Younes, dirigeant d'Alim Express (PACA)"],
         ["Réalisation","Yassin (cœur back-office) · David (API, clients, facturation, livraisons)"],
         ["Stack","Next.js 15 · PostgreSQL · Prisma · NextAuth · Electron"],
         ["Période","3 mars → 22 juin 2026"]])

    b.callout([("Périmètre de ce dossier. ",True),
        ("Ce document couvre uniquement le BACK-OFFICE d'Alim Express : l'outil d'administration "
         "interne utilisé par l'équipe (produits, stock, clients, commandes, facturation, livraisons) "
         "et l'API REST qu'il expose au portail client. Le portail client fait l'objet d'un dossier "
         "technique distinct. Organisation selon les cinq critères de la grille « Réalisation du projet "
         "— Développeur full stack ».",False)])

    b.h("Index des livrables",2)
    b.table(["Livrable","Type","Accès"],[
        ["Code source back-office","GitHub","github.com/mineboxarabic/AlimExpressApp"],
        ["Modèle de données","Prisma","prisma/schema.prisma (~980 lignes, 25 migrations)"],
        ["API REST consommée par le portail","Code","src/app/api/v1/*"],
        ["Tableau de pilotage","Trello","lien partagé"],
        ["Spécifications & conventions","Markdown","AGENTS.md, README.md"],
        ["Installeur desktop","Electron","AlimExpressApp-Setup-{version}.exe"],
    ])

    # ---- Critère 1 ----
    b.crit("/ 5","Critère 1 — Structuration technique de l'application")
    b.h("1.1 Architecture — Feature-Sliced Design (FSD)",3)
    b.p("Le back-office applique une architecture en couches (Feature-Sliced Design) adaptée à l'App "
        "Router de Next.js : la couche de routage reste fine, la logique métier est isolée par domaine "
        "fonctionnel (21 features), et les utilitaires/UI sont mutualisés.")
    b.code("""src/
├── app/            routage (dashboard protégé, auth, api/)
├── features/       21 domaines métier : products, inventory, orders, invoices,
│                   credit-notes, payments, deliveries, palettes, suppliers,
│                   customers, retail, cart, catalog, dashboard, order-history,
│                   users, api-keys, ai-assistant, system, auth, order-validation
│      └── chaque feature : *.action.ts · *.schema.ts · components/ · utils/
├── lib/            db, email, whatsapp, storage/s3, pdf/, geo/, translations…
├── components/     ui/ (shadcn) + crud-layout réutilisable
├── widgets/        sidebar, header, title-bar Electron, health-monitor
├── contexts/  ·  types/""")

    b.h("1.2 Stack technique",3)
    b.table(["Domaine","Technologie","Version"],[
        ["Framework","Next.js (App Router, output standalone)","15.3.9"],
        ["UI","React / TypeScript strict","19.2.3 / 5.x"],
        ["Base de données","PostgreSQL + Prisma ORM","6.2"],
        ["Auth admin","NextAuth (session, credentials)","5 (beta)"],
        ["Auth clients","jose (JWT B2B & B2C, HS256)","6.2.3"],
        ["UI / style","Tailwind CSS + shadcn/ui (Radix)","3.4 / —"],
        ["PDF / codes-barres","jsPDF + jspdf-autotable + JSBarcode","4.2 / 5.0 / 3.12"],
        ["E-mail","Resend + Nodemailer (SMTP secours)","6.5 / 7.0"],
        ["WhatsApp","Twilio","5.10"],
        ["Stockage","AWS S3 (SDK v3)","3.x"],
        ["IA","Google Gemini, OpenRouter, OpenAI + MCP SDK","—"],
        ["Graphiques / DnD","recharts · dnd-kit","2.15 / 6.x"],
        ["Desktop","Electron + electron-builder (NSIS)","42 / 26"],
        ["Tests","Vitest + Playwright + Testing Library","4.0 / 1.57"],
    ])

    b.h("1.3 Modèle de données (Prisma)",3)
    b.p("Schéma PostgreSQL de ~980 lignes, 25 migrations versionnées, ~25 modèles couvrant tout le "
        "domaine métier. Entités principales :")
    b.table(["Domaine","Entités principales"],[
        ["Catalogue","Product (prix C/D/E/F, visibilité), Category, Marque, Tag (+ tables de liaison)"],
        ["Stock","InventoryItem (lot, péremption, batch), StockMovement (audit), Palette, PaletteItem, LotAllocation (FEFO)"],
        ["Clients","Customer (B2B : SIRET, TVA, APE, niveau tarifaire, géoloc), RetailCustomer (B2C)"],
        ["Commercial","Order / OrderItem (B2B), RetailOrder / RetailOrderItem (B2C), Quote (devis signé)"],
        ["Facturation","Invoice, InvoiceItem, InvoiceVersion, InvoiceStatusHistory, Payment, CreditNote (avoir)"],
        ["Logistique","Delivery (tournée), DeliveryCity (arrêt ordonné, géoloc, verrouillage)"],
        ["Système","User (rôles), ApiKey, AuthToken, SystemSettings, OrderHistoryEvent (audit), ChatSession/Message (IA)"],
    ])
    b.p("Énumérations notables : Role (ADMIN/MANAGER/STAFF/B2B_CLIENT/B2C_CLIENT), PricingLevel (C/D/E/F), "
        "ProductVisibility (RETAIL/WHOLESALE/BOTH), OrderStatus, InvoiceStatus, StockMovementType, "
        "LotStatus, DeliveryStatus.", color=MUTED, size=9.5)
    b.ph("diagram","Schéma BO-1 — Architecture en couches (FSD)",
        "app (routage) → features (21 domaines) → lib (services partagés) → components/widgets, avec l'API exposée.",
        "→ Prompt fourni séparément (Schéma BO-1)")
    b.ph("diagram","Schéma BO-2 — Modèle de données (entités & relations)",
        "Diagramme entité-relation des modèles principaux (Product, InventoryItem, Order, Quote, Invoice, CreditNote, Delivery…).",
        "→ Prompt fourni séparément (Schéma BO-2)")

    b.h("1.4 Authentification — trois publics + clés API",3)
    b.bullets([
        ["Administrateurs / staff : ","NextAuth v5 (session, cookie HttpOnly, mot de passe haché bcrypt), rôles ADMIN/MANAGER/STAFF."],
        ["Clients B2B & B2C : ","JWT signés (jose, HS256, secret AUTH_SECRET, validité 30 j) — vérifiables par le portail."],
        ["Intégrations externes (MCP) : ","clés API hachées HMAC-SHA256, en-tête Authorization/x-api-key."],
    ])

    b.h("1.5 API REST exposée (/api/v1)",3)
    b.p("Le back-office héberge l'API consommée par le portail. Familles d'endpoints :")
    b.table(["Famille","Endpoints (extraits)"],[
        ["Auth","/api/v1/auth/{verify-email, set-password, forgot-password, reset-password} ; /api/auth/[...nextauth]"],
        ["B2B","/api/v1/b2b/{auth/login, me, catalog, cart, orders, quotes, quotes/[id]/accept, quotes/[id]/pdf, invoices/[id]/pdf}"],
        ["B2C (retail)","/api/v1/retail/{auth/login, auth/register, me, catalog, cart, orders/checkout, orders/me/[id]}"],
        ["Public","/api/v1/public/{categories, marques, deliveries/upcoming}"],
        ["Admin (CRUD)","/api/v1/{products, customers, orders, invoices, api-keys}"],
        ["Automatisé","/api/cron/lock-deliveries (verrouillage J-1) · /api/mcp (IA) · /api/ai/chat"],
    ])

    b.h("1.6 Points techniques notables & qualité",3)
    b.bullets([
        ["Gestion FEFO du stock : ","LotAllocation relie chaque ligne de commande au lot le plus proche de la péremption ; StockMovement trace tous les mouvements."],
        ["Verrouillage J-1 21h : ","tâche planifiée (cron) qui fige les commandes de la tournée du lendemain (lockedAt)."],
        ["Facturation versionnée : ","InvoiceVersion + InvoiceStatusHistory + figeage (snapshot) de la facture à l'envoi."],
        ["Génération documentaire : ","devis, factures, avoirs, bons de livraison/chargement en PDF (jsPDF + codes-barres), stockés sur S3."],
        ["Assistant IA (MCP) : ","exposition de 50+ outils métier via le protocole MCP pour des agents externes."],
        ["Qualité : ","SOLID, composants ≤ 200 lignes, Server Actions priorisées sur les routes API, Zod (source unique de validation), aucun type any, démarche TDD."],
    ])

    # ---- Critère 2 ----
    b.crit("/ 3","Critère 2 — Contraintes d'hébergement anticipées")
    b.p("Le back-office est conteneurisé et déployable en cloud ou en poste local (desktop). Les choix "
        "privilégient la portabilité et la sobriété tarifaire.")
    b.table(["Élément","Choix"],[
        ["Conteneurisation","Dockerfile multi-étapes (Node 22 Alpine) + docker-compose (PostgreSQL 15)"],
        ["Hébergement cloud","Railway / Vercel — sortie standalone Next.js"],
        ["Tâche planifiée","vercel.json : /api/cron/lock-deliveries à 19h & 20h UTC (gère l'heure d'été), protégée par CRON_SECRET"],
        ["Stockage fichiers","AWS S3 (images produit + PDF documents)"],
        ["Mode desktop","Electron (Windows NSIS) avec base SQLite locale — usage hors-ligne"],
    ])
    b.h("Variables d'environnement (extrait de .env.example)",4)
    b.bullets([
        ["Base / Auth : ","DATABASE_URL, AUTH_SECRET (JWT partagé), NEXTAUTH_SECRET/URL, CRON_SECRET."],
        ["E-mail : ","RESEND_API_KEY, EMAIL_FROM, SMTP_* (secours)."],
        ["WhatsApp : ","TWILIO_ACCOUNT_SID, TWILIO_AUTH_TOKEN, TWILIO_WHATSAPP_FROM."],
        ["S3 : ","AWS_REGION (eu-west-3), AWS_ACCESS_KEY_ID/SECRET, S3_BUCKET_NAME."],
        ["Entreprise / portail : ","COMPANY_* (figées sur les documents), CLIENT_PORTAL_URL, ALLOWED_ORIGINS (CORS)."],
        ["IA : ","OPENAI_API_KEY, OPENAI_IMAGE_MODEL, GOOGLE_GENERATIVE_AI_API_KEY."],
    ])

    # ---- Critère 3 ----
    b.crit("/ 5","Critère 3 — Administration centralisée (le back-office)")
    b.p("Le back-office EST l'outil d'administration : il centralise toute la gestion d'Alim Express "
        "dans une interface unique, pour l'équipe interne.")
    b.ph("capture","Capture BO-1 — Tableau de bord",
        "Vue d'ensemble : commandes, statuts de factures, encours, graphique des ventes (recharts).",
        "→ Source fournie séparément (Capture BO-1)")
    b.table(["Module","Fonctions"],[
        ["Produits","CRUD, SKU/codes-barres, images (S3, amélioration IA), niveaux de prix C/D/E/F, visibilité, brouillons."],
        ["Stock","Lots (péremption, FEFO), palettes (réception DRAFT→CONFIRMED), mouvements tracés, statuts ACTIVE/QUARANTINE."],
        ["Clients","B2B (société, SIRET, TVA, APE, niveau tarifaire, géoloc) et B2C ; création des comptes portail."],
        ["Commercial","Commandes B2B & retail, génération de devis (PDF), suivi des statuts."],
        ["Facturation","Factures (cycle brouillon→envoyée→payée), versions & historique de statut, avoirs, paiements (partiels)."],
        ["Livraisons","Tournées multi-villes ordonnées, géolocalisation, verrouillage J-1 21h, bons de livraison/chargement PDF."],
        ["Pilotage","Tableau de bord & analytics (recharts), journal d'audit (OrderHistoryEvent)."],
        ["Système","Utilisateurs & rôles, clés API, paramètres (e-mail/Twilio/S3), multilingue FR/AR, assistant IA."],
    ])
    b.ph("capture","Capture BO-2 — Gestion des produits",
        "Liste des produits + fiche produit montrant les niveaux de prix C/D/E/F et la visibilité.",
        "→ Source fournie séparément (Capture BO-2)")
    b.ph("capture","Capture BO-3 — Stock / Livraisons",
        "Écran de gestion des lots & palettes (FEFO) OU planning des tournées de livraison.",
        "→ Source fournie séparément (Capture BO-3)")
    b.ph("capture","Capture BO-4 — Facturation",
        "Édition d'une facture (et/ou PDF généré : facture / devis / avoir).",
        "→ Source fournie séparément (Capture BO-4)")
    b.h("Spécifications techniques et fonctionnelles écrites",4)
    b.bullets([
        ["Modèle de données : ","prisma/schema.prisma (~980 lignes), 25 migrations, déployable sur base vierge."],
        ["Contrat d'API : ","endpoints /api/v1/* avec validation Zod en entrée/sortie."],
        ["Conventions : ","AGENTS.md (FSD, SOLID, TDD, nommage) + README.md."],
    ])
    b.ph("diagram","Schéma BO-3 — Cycle commercial & impact stock",
        "Devis → acceptation signée → commande ferme → livraison (déduction FEFO) → facture → avoir (remise en stock).",
        "→ Prompt fourni séparément (Schéma BO-3)")

    # ---- Critère 4 ----
    b.crit("/ 4","Critère 4 — Recettage formalisé et corrections")
    b.p("Démarche TDD obligatoire (Red → Green → Refactor). Trois niveaux de tests automatisés + "
        "recette manuelle avec le commanditaire (procès-verbal signé).")
    b.bullets([
        ["Tests unitaires : ","Vitest (jsdom) + Testing Library, co-localisés (*.test.ts)."],
        ["Tests d'intégration : ","Vitest sur base PostgreSQL isolée (docker-compose.test.yml, *.integration.test.ts)."],
        ["Tests end-to-end : ","Playwright — auth, workflow critique commande→devis→facture→paiement, dashboard, génération PDF, produits."],
        ["Couverture cible : ","> 60 % sur les modules critiques."],
    ])
    b.ph("capture","Capture BO-5 — Exécution des tests",
        "Sortie Vitest / Playwright au vert et/ou rapport de couverture (test:coverage).",
        "→ Source fournie séparément (Capture BO-5)")
    b.code("""npm run test            # tests unitaires (Vitest)
npm run test:coverage   # couverture
npm run test:integration# tests d'intégration (base de test Docker)
npm run test:e2e        # tests end-to-end (Playwright)""")

    # ---- Critère 5 ----
    b.crit("/ 3","Critère 5 — Accessibilité (responsive + desktop)")
    b.p("Interface d'administration responsive (barre latérale repliable, tiroirs sur mobile) et "
        "multilingue FR/AR. En complément, l'application est packagée en version desktop (Electron, "
        "Windows) avec base locale pour un usage hors-ligne au comptoir/entrepôt.")
    b.ph("capture","Capture BO-6 — Desktop / Assistant IA",
        "Application desktop Electron (barre de titre personnalisée) et/ou assistant IA intégré.",
        "→ Source fournie séparément (Capture BO-6)")
    b.h("Repère — charte graphique",4)
    b.table(["Usage","Hex"],[
        ["Vert principal","#3F561F"],["Fond","#FAF8F2"],["Section","#F0EBDD"],
        ["Texte","#171717"],["Alerte","#D52B14"],["Mise en avant","#F2C400"]])

    b.h("Annexe — Glossaire",2)
    b.table(["Terme","Définition"],[
        ["FEFO","First-Expired-First-Out : sortie du lot dont la péremption est la plus proche."],
        ["Devis / Avoir","Proforma soumis à acceptation / document de crédit avec remise en stock."],
        ["MCP","Model Context Protocol : exposition d'outils métier à des agents IA."],
        ["FSD","Feature-Sliced Design : architecture en couches par domaine."],
        ["Niveau tarifaire","Catégorie de prix C/D/E/F attribuée à un client professionnel."],
    ])
    b.save(r"c:\Users\anasl\Documents\AlimExpressCatalog\docs\doc\Dossier-technique-BackOffice-AlimExpress.docx")

# ===================================================================================
# DOSSIER 2 — PORTAIL CLIENT (AlimExpressCatalog) — plus détaillé
# ===================================================================================
def build_portail():
    b=Doc()
    b.cover("MBA DÉVELOPPEUR FULL STACK 2 · MYDIGITALSCHOOL · WORKSHOP CLIENT",
        "Dossier technique — Portail client",
        "Alim Express — storefront e-commerce B2C (particuliers) et B2B (professionnels)",
        "Enseigne · Le Bakkal Oriental",
        [["Application","AlimExpressCatalog (portail client)"],
         ["Commanditaire","Rashed Younes, dirigeant d'Alim Express (PACA)"],
         ["Réalisation","Anas (portail intégral) · David (API de liaison)"],
         ["Stack","Next.js 15 · TypeScript · Tailwind v4 · jose · Zod"],
         ["Nature","Client de l'API back-office — sans accès direct à la base"]])

    b.callout([("Périmètre de ce dossier. ",True),
        ("Ce document couvre uniquement le PORTAIL CLIENT d'Alim Express : la vitrine en ligne "
         "destinée aux particuliers (B2C, « retail ») et aux professionnels (B2B, « pro »). Le portail "
         "est un client de l'API du back-office (dossier distinct) et n'accède pas directement à la base "
         "de données. Organisation selon les cinq critères de la grille « Réalisation du projet ».",False)])

    b.h("Index des livrables",2)
    b.table(["Livrable","Type","Accès"],[
        ["Code source portail","GitHub","github.com/mineboxarabic/albakkaloriental_front"],
        ["Environnement de démonstration","App déployée","URL de démo"],
        ["Documentation de tests","Markdown","docs/doc/TESTING.md"],
        ["Conventions","Markdown","AGENTS.md, README.md"],
        ["Tableau de pilotage","Trello","lien partagé"],
    ])

    # ---- Critère 1 ----
    b.crit("/ 5","Critère 1 — Structuration technique de l'application")
    b.h("1.1 Architecture — client d'API, groupes de routes",3)
    b.p("Le portail est une application Next.js 15 (App Router) qui consomme l'API REST du back-office. "
        "Il n'accède jamais directement à la base : le back-office est l'unique source de vérité. "
        "Les parcours sont isolés par groupes de routes.")
    b.code("""app/
├── (public)/   CGV, confidentialité, contact, mots de passe, vérification e-mail
├── (retail)/   B2C : accueil, products, products/[id], cart, checkout,
│               account, orders, orders/[id], categories, marques, login, register
└── (pro)/      B2B : pro/login, pro/products(+[id]), pro/cart, pro/quotes(+[id]),
                pro/orders(+[id]), pro/invoices, pro/account  (+ route handlers PDF)
actions/        Server Actions validées par Zod (auth, cart, order, quote, me…)
lib/            api-client, session, catalog, catalog-pricing, catalog-sort, order-rules…
middleware.ts   protection des routes /pro/* (vérification JWT)""")

    b.h("1.2 Stack technique",3)
    b.table(["Domaine","Technologie","Version"],[
        ["Framework","Next.js (App Router)","15.5.18"],
        ["UI","React / TypeScript strict","19.2.4 / 5.x"],
        ["Style","Tailwind CSS + shadcn/ui + lucide","4.x"],
        ["Authentification","jose (vérification JWT côté portail)","6.2.3"],
        ["Validation","Zod","4.4.3"],
        ["Dates","date-fns","4.2.1"],
        ["Tests unitaires","Vitest + Testing Library (jsdom)","4.1.6"],
        ["Tests E2E","Playwright (Chromium)","1.60"],
    ])

    b.h("1.3 Server Actions (mutations validées Zod)",3)
    b.table(["Fichier","Rôle"],[
        ["retail-auth.ts / pro-auth.ts","Inscription, connexion, déconnexion (stockage du JWT en cookie)."],
        ["retail-me.ts / pro-me.ts","Profil, mise à jour, changement de mot de passe, commandes, factures."],
        ["retail-cart.ts / pro-cart.ts","Panier serveur : ajout, mise à jour, suppression, vidage."],
        ["retail-order.ts / pro-order.ts","Validation de commande (checkout)."],
        ["pro-quote.ts","Devis : liste, détail, acceptation signée (transmet IP/User-Agent)."],
        ["password-reset.ts / set-password.ts / verify-email.ts","Flux mots de passe & vérification e-mail."],
    ])

    b.h("1.4 Utilitaires lib/ (extraits)",3)
    b.bullets([
        ["api-client.ts : ","client REST (backendFetch) — injection du Bearer depuis le cookie, gestion d'erreurs, déconnexion auto sur 401, cache no-store."],
        ["session.ts : ","vérification du JWT (jose), cookie catalog_session (HttpOnly, SameSite=Lax, Secure en prod, 30 j)."],
        ["catalog-pricing.ts : ","calcul du prix par niveau (getTierPrice C/D/E/F, repli sellingPrice), unités PACK/UNIT."],
        ["order-rules.ts : ","règles métier centralisées (valeurs exactes ci-dessous)."],
        ["catalog.ts / catalog-sort.ts : ","récupération du catalogue par audience, filtres visibilité, tri (prix/nom, en stock d'abord)."],
    ])
    b.table(["Règle métier (retail)","Valeur"],[
        ["Montant minimum de commande","30 €"],
        ["Franco de port (livraison offerte)","≥ 50 €"],
        ["Frais de livraison sinon","3 €"],
        ["Quantité max par produit (B2C)","3"],
        ["Quantité par produit (B2B)","non plafonnée"],
    ])

    b.h("1.5 Authentification & sécurité",3)
    b.bullets([
        ["Cookie de session : ","catalog_session — JWT émis par le back-office, HttpOnly, SameSite=Lax, Secure (prod), durée 30 jours."],
        ["Middleware : ","protège /pro/* (panier, devis, commandes, factures, compte) ; vérifie le JWT (jose) et le rôle (B2B_CLIENT/ADMIN) ; redirige vers /pro/login?next=…"],
        ["Catalogue pro public : ","/pro/products, /pro/categories, /pro/marques restent consultables sans prix ; les prix négociés ne sont jamais exposés à un visiteur non authentifié."],
        ["Rôles JWT : ","B2C_CLIENT (retail) · B2B_CLIENT / ADMIN (pro)."],
    ])
    b.ph("diagram","Schéma PC-1 — Architecture du portail",
        "Portail client (groupes (public)/(retail)/(pro)) → API REST back-office (JWT) ; aucun accès direct à la base.",
        "→ Prompt fourni séparément (Schéma PC-1)")

    # ---- Parcours détaillés (le « plus de détails » demandé) ----
    b.h("Parcours détaillé — Client particulier (B2C)",2)
    b.bullets([
        ["1. Découverte : ","accueil (héro, prochaines livraisons, catégories, meilleures ventes, nouveautés)."],
        ["2. Navigation : ","/products avec recherche (nom/SKU), tri (prix, nom), filtres catégorie & marque."],
        ["3. Fiche produit : ","/products/[id] — prix de détail, unité de vente ; l'ajout au panier invité ouvre la modale de connexion."],
        ["4. Inscription : ","/register/particulier (coordonnées, adresse, mot de passe, acceptation CGV/RGPD) → vérification e-mail."],
        ["5. Panier : ","révision, modification des quantités (plafonnées à 3), suppression ; calcul du franco de port."],
        ["6. Commande : ","/checkout (adresse de livraison pré-remplie, notes) → POST checkout → confirmation."],
        ["7. Suivi : ","/orders et /orders/[id] (statut PENDING→…→DELIVERED), /account (profil, mot de passe)."],
    ])
    b.ph("diagram","Schéma PC-2 — Parcours B2C",
        "Découverte → catalogue → fiche → (connexion) → panier → checkout → confirmation → suivi de commande.",
        "→ Prompt fourni séparément (Schéma PC-2)")
    b.ph("capture","Capture PC-1 — Accueil / Catalogue B2C",
        "Page d'accueil (héro + catégories) ou listing /products avec recherche et tri.",
        "→ Source fournie séparément (Capture PC-1)")
    b.ph("capture","Capture PC-2 — Fiche produit & panier B2C",
        "Fiche produit /products/[id] et panier /cart (quantités, franco de port).",
        "→ Source fournie séparément (Capture PC-2)")

    b.h("Parcours détaillé — Client professionnel (B2B)",2)
    b.p("Il n'y a pas d'auto-inscription B2B : la qualification commerciale (validation du KBIS) est "
        "préalable à toute création de compte, et les identifiants sont créés manuellement depuis le "
        "back-office.")
    b.bullets([
        ["1. Demande de compte : ","le professionnel remplit le formulaire société sur /register/entreprise (raison sociale, SIRET, contact, adresse)."],
        ["2. Envoi du KBIS via WhatsApp : ","le bouton « Envoyer mon KBIS via WhatsApp » ouvre un message pré-rempli vers le +33 7 66 30 13 39 ; le pro y joint son KBIS (PDF ou photo) et l'envoie."],
        ["3. Validation & création du compte : ","l'équipe vérifie le dossier et crée les identifiants professionnels manuellement depuis le back-office (réponse sous 24 à 48 h)."],
        ["4. Activation : ","le pro reçoit un lien e-mail pour définir son mot de passe."],
        ["5. Connexion : ","/pro/login (e-mail + mot de passe), redirection vers la cible demandée."],
        ["6. Catalogue négocié : ","/pro/products — prix selon le niveau tarifaire (C/D/E/F) du client, unités PACK/UNIT, recherche/tri/filtres."],
        ["7. Panier : ","/pro/cart — quantités non plafonnées (usage grossiste)."],
        ["8. Devis (proforma) : ","/pro/quotes — génération, détail (lignes, TVA, total, date de validité)."],
        ["9. Acceptation signée : ","acceptation horodatée (IP + User-Agent enregistrés) → commande ferme (CONFIRMED)."],
        ["10. Suivi & documents : ","/pro/orders, /pro/invoices (statuts UNPAID/PAID/PARTIAL/OVERDUE), téléchargement des PDF devis & factures, /pro/account (société, SIRET, TVA, niveau, encours)."],
    ])
    b.ph("diagram","Schéma PC-3 — Parcours B2B (de la demande à la facture)",
        "Demande /register/entreprise → envoi du KBIS via WhatsApp → validation & création du compte au back-office → activation par lien → connexion → catalogue prix négociés → panier → devis → acceptation signée → commande ferme → facture (PDF).",
        "→ Prompt fourni séparément (Schéma PC-3)")
    b.ph("capture","Capture PC-3 — Demande de compte pro (KBIS via WhatsApp)",
        "Formulaire /register/entreprise avec le bouton « Envoyer mon KBIS via WhatsApp » et les étapes 1–4.",
        "→ Source fournie séparément (Capture PC-3)")
    b.ph("capture","Capture PC-4 — Espace pro (catalogue & prix par niveau)",
        "Connexion pro + catalogue /pro/products affichant les prix selon le niveau tarifaire.",
        "→ Source fournie séparément (Capture PC-4)")
    b.ph("capture","Capture PC-5 — Devis & factures pro",
        "Détail d'un devis (/pro/quotes/[id]) + liste des factures avec statuts et téléchargement PDF.",
        "→ Source fournie séparément (Capture PC-5)")

    # ---- Critère 2 ----
    b.crit("/ 3","Critère 2 — Contraintes d'hébergement anticipées")
    b.p("Application Next.js sans état (session portée par le cookie JWT), donc simple à héberger et à "
        "redimensionner ; aucune base de données propre.")
    b.table(["Élément","Choix"],[
        ["Hébergement","Plateforme compatible Next.js (Railway / Vercel) — build & déploiement continus depuis Git"],
        ["Images produit","AWS S3 (alimexpress-buckeyt, région eu-north-1), domaine autorisé dans next.config.ts"],
        ["Variables d'environnement","BACKEND_URL, NEXT_PUBLIC_BACKEND_URL, AUTH_SECRET (secret JWT partagé)"],
        ["État","Sans état côté serveur — session dans le cookie HttpOnly catalog_session"],
        ["Base de données","Aucune — client d'API pur (Prisma/migrations interdits côté portail)"],
    ])

    # ---- Critère 3 ----
    b.crit("/ 5","Critère 3 — Administration via le back-office")
    b.p("Le portail ne contient pas d'administration propre : son contenu et ses comptes sont pilotés "
        "depuis le back-office AlimExpressApp (dossier distinct). Cette séparation isole la surface "
        "publique de la surface interne.")
    b.bullets([
        ["Catalogue & prix : ","produits, niveaux tarifaires C/D/E/F, visibilité RETAIL/WHOLESALE/BOTH gérés au back-office."],
        ["Comptes pro : ","demande via /register/entreprise + envoi du KBIS par WhatsApp ; l'équipe valide le dossier et crée les identifiants manuellement au back-office (activation par lien) — pas d'auto-inscription B2B."],
        ["Cycle commercial : ","devis, commandes, factures et avoirs traités au back-office ; le portail en restitue l'état."],
        ["Spécifications : ","contrat d'API /api/v1/* documenté + AGENTS.md/README.md du portail."],
    ])

    # ---- Critère 4 ----
    b.crit("/ 4","Critère 4 — Recettage formalisé et corrections")
    b.p("Deux niveaux de tests automatisés + recette manuelle.")
    b.bullets([
        ["Tests unitaires / composants : ","Vitest + Testing Library — dossiers __tests__/actions, __tests__/components, __tests__/lib (logique de prix, tri, règles, rendu)."],
        ["Tests end-to-end : ","Playwright (ports 3090/3091) contre un back-end mocké en mémoire (e2e/mock-backend.ts)."],
    ])
    b.h("Scénarios E2E couverts (e2e/cart.spec.ts)",4)
    b.table(["Scénario","Vérification"],[
        ["B2C — ajout invité","Connexion requise → modale → login → reprise de l'intention d'ajout."],
        ["B2C — plafond quantité","Quantité bloquée à 3 par produit (bouton + désactivé)."],
        ["B2C — session expirée","Cookie invalidé → alerte « session expirée » → redirection /login."],
        ["B2B — grosses quantités","Ajout > 3 sans plafonnement (panier pro), montée jusqu'à 5–6."],
    ])
    b.ph("capture","Capture PC-6 — Exécution des tests",
        "Sortie Vitest + Playwright au vert (et/ou rapport de couverture).",
        "→ Source fournie séparément (Capture PC-6)")
    b.code("""npx vitest run --pool=threads      # tests unitaires et composants
npx playwright test                 # tests end-to-end
npm run lint                        # qualité du code (ESLint)""")
    b.p("Corrections réalisées (traçables dans Git) : refonte de la navigation retail suite aux retours "
        "client, ajustements responsive, intégration du branding Le Bakkal Oriental (logo, favicon).")

    # ---- Critère 5 ----
    b.crit("/ 3","Critère 5 — Responsive et mobile-first")
    b.p("Le portail est conçu en approche mobile-first : pensé et validé prioritairement sur smartphone "
        "(usage réel des clients), puis décliné sur tablette et ordinateur. Grilles fluides Tailwind, "
        "composants adaptatifs (menu, panier coulissant, cartes produit).")
    b.ph("capture","Capture PC-7 — Responsive mobile / desktop",
        "Même page catalogue en vue mobile (~390 px) et en vue desktop, côte à côte.",
        "→ Source fournie séparément (Capture PC-7)")
    b.h("Charte graphique — Le Bakkal Oriental",4)
    b.table(["Usage","Couleur","Hex"],[
        ["Vert principal (CTA, nav)","Olive foncé","#3F561F"],
        ["Fond principal","Blanc cassé chaud","#FAF8F2"],
        ["Fond de section","Beige crème","#F0EBDD"],
        ["Texte principal","Noir","#171717"],
        ["Texte secondaire","Gris chaud","#6B665D"],
        ["Bordures","Beige-gris clair","#DDD8CC"],
        ["Alerte / promo","Rouge-orangé","#D52B14"],
        ["Mise en avant","Jaune chaud","#F2C400"],
    ])
    b.p("Typographie : Poppins (titres) et Satoshi (corps).")

    b.h("Annexe — Glossaire",2)
    b.table(["Terme","Définition"],[
        ["B2C / B2B","Particuliers (prix standards) / professionnels (prix négociés par niveau)."],
        ["Devis (proforma)","Document chiffré pré-commande, soumis au client pour acceptation signée."],
        ["Niveau tarifaire","Catégorie de prix C/D/E/F attribuée à un client professionnel."],
        ["catalog_session","Cookie HttpOnly portant le JWT de session du client."],
        ["Franco de port","Seuil (50 €) au-delà duquel la livraison est offerte."],
    ])
    b.save(r"c:\Users\anasl\Documents\AlimExpressCatalog\docs\doc\Dossier-technique-Portail-Client-AlimExpress.docx")

build_backoffice()
build_portail()
