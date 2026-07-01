# -*- coding: utf-8 -*-
"""Fusionne les dossiers back-office + portail (avec leurs images) en un seul document,
ajoute la page de garde commune (accès Git + liens apps) et les annexes vides."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docxcompose.composer import Composer

GREEN=RGBColor(0x3F,0x56,0x1F); INK=RGBColor(0x17,0x17,0x17); MUTED=RGBColor(0x6B,0x66,0x5D)
WHITE=RGBColor(0xFF,0xFF,0xFF)
GREEN_HEX="3F561F"; SECTION_HEX="F0EBDD"; HILITE_HEX="F2C400"; GREENSOFT_HEX="EAEDE3"; FONT="Poppins"
DOCDIR=r"c:\Users\anasl\Documents\AlimExpressCatalog\docs\doc"

def _shade(el,hx):
    sh=OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:color"),"auto")
    sh.set(qn("w:fill"),hx); el.append(sh)
def _borders(pPr,hx="DDD8CC",sz="8",space="6",sides=("top","left","bottom","right")):
    pbdr=OxmlElement("w:pBdr")
    for s in sides:
        b=OxmlElement(f"w:{s}"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),sz)
        b.set(qn("w:space"),space); b.set(qn("w:color"),hx); pbdr.append(b)
    pPr.append(pbdr)

class Doc:
    def __init__(self):
        self.d=Document()
        n=self.d.styles["Normal"]; n.font.name=FONT; n.font.size=Pt(10.5)
        n.font.color.rgb=INK; n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.22
    def run(self,p,t,*,size=10.5,bold=False,color=INK,italic=False):
        r=p.add_run(t); r.font.name=FONT; r.font.size=Pt(size); r.bold=bold; r.italic=italic
        r.font.color.rgb=color; return r
    def h(self,t,level=2):
        p=self.d.add_paragraph(); p.paragraph_format.space_before=Pt(16 if level==2 else 10)
        p.paragraph_format.space_after=Pt(6); sizes={1:20,2:15,3:12.5,4:11}
        self.run(p,t,size=sizes[level],bold=True,color=GREEN)
        if level==2: _borders(p._p.get_or_add_pPr(),HILITE_HEX,"18","4",("bottom",))
        return p
    def p(self,t=None,*,size=10.5,color=INK,bold=False,italic=False,sa=6):
        p=self.d.add_paragraph(); p.paragraph_format.space_after=Pt(sa)
        if t is not None: self.run(p,t,size=size,bold=bold,color=color,italic=italic)
        return p
    def callout(self,parts):
        p=self.d.add_paragraph(); pPr=p._p.get_or_add_pPr(); _shade(pPr,SECTION_HEX)
        _borders(pPr,GREEN_HEX,"24","8",("left",))
        p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(10)
        p.paragraph_format.left_indent=Pt(10); p.paragraph_format.right_indent=Pt(6)
        for txt,bold in parts: self.run(p,txt,color=(GREEN if bold else MUTED),bold=bold,size=10)
    def hyperlink(self,p,url,text):
        r_id=p.part.relate_to(url,RT.HYPERLINK,is_external=True)
        link=OxmlElement("w:hyperlink"); link.set(qn("r:id"),r_id)
        r=OxmlElement("w:r"); rPr=OxmlElement("w:rPr")
        c=OxmlElement("w:color"); c.set(qn("w:val"),GREEN_HEX); rPr.append(c)
        u=OxmlElement("w:u"); u.set(qn("w:val"),"single"); rPr.append(u)
        rf=OxmlElement("w:rFonts"); rf.set(qn("w:ascii"),FONT); rf.set(qn("w:hAnsi"),FONT); rPr.append(rf)
        sz=OxmlElement("w:sz"); sz.set(qn("w:val"),"19"); rPr.append(sz)
        r.append(rPr); t=OxmlElement("w:t"); t.text=text; r.append(t); link.append(r); p._p.append(link)
    def link_table(self,headers,rows):
        t=self.d.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
        for i,htext in enumerate(headers):
            c=t.rows[0].cells[i]; _shade(c._tc.get_or_add_tcPr(),GREEN_HEX)
            pp=c.paragraphs[0]; pp.paragraph_format.space_after=Pt(2)
            self.run(pp,htext,bold=True,color=WHITE,size=9.5)
        for ri,(label,url) in enumerate(rows):
            cells=t.add_row().cells
            if ri%2==1:
                _shade(cells[0]._tc.get_or_add_tcPr(),GREENSOFT_HEX); _shade(cells[1]._tc.get_or_add_tcPr(),GREENSOFT_HEX)
            p0=cells[0].paragraphs[0]; p0.paragraph_format.space_after=Pt(2)
            self.run(p0,label,size=9.5,bold=True,color=GREEN)
            p1=cells[1].paragraphs[0]; p1.paragraph_format.space_after=Pt(2)
            self.hyperlink(p1,url,url)
        self.d.add_paragraph().paragraph_format.space_after=Pt(2)
    def kv_table(self,rows):
        t=self.d.add_table(rows=1,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
        for i,htext in enumerate(["Information","Valeur"]):
            c=t.rows[0].cells[i]; _shade(c._tc.get_or_add_tcPr(),GREEN_HEX)
            self.run(c.paragraphs[0],htext,bold=True,color=WHITE,size=9.5)
        for ri,(k,v) in enumerate(rows):
            cells=t.add_row().cells
            if ri%2==1:
                _shade(cells[0]._tc.get_or_add_tcPr(),GREENSOFT_HEX); _shade(cells[1]._tc.get_or_add_tcPr(),GREENSOFT_HEX)
            self.run(cells[0].paragraphs[0],k,size=9.3,bold=True,color=GREEN)
            self.run(cells[1].paragraphs[0],v,size=9.3)
        self.d.add_paragraph().paragraph_format.space_after=Pt(2)
    def empty_box(self,hint,lines=10):
        p=self.d.add_paragraph(); pPr=p._p.get_or_add_pPr(); _shade(pPr,"FFFFFF")
        _borders(pPr,GREEN_HEX,"10","10")
        p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6)
        p.paragraph_format.left_indent=Pt(8); p.paragraph_format.right_indent=Pt(8)
        self.run(p,hint+("\n"*lines),size=9,italic=True,color=MUTED)
    def page_break(self): self.d.add_page_break()
    def save(self,path): self.d.save(path)

# ---------- INTRO (page de garde commune + accès) ----------
intro=Doc()
for _ in range(2): intro.d.add_paragraph()
p=intro.d.add_paragraph(); _shade(p._p.get_or_add_pPr(),GREEN_HEX)
p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(2); p.paragraph_format.left_indent=Pt(8)
intro.run(p,"MBA DÉVELOPPEUR FULL STACK 2 · MYDIGITALSCHOOL · WORKSHOP CLIENT",bold=True,size=9,color=RGBColor(0xF2,0xC4,0x00))
p2=intro.d.add_paragraph(); _shade(p2._p.get_or_add_pPr(),GREEN_HEX); p2.paragraph_format.left_indent=Pt(8); p2.paragraph_format.space_after=Pt(2)
intro.run(p2,"Dossier technique — Alim Express",bold=True,size=34,color=WHITE)
p3=intro.d.add_paragraph(); _shade(p3._p.get_or_add_pPr(),GREEN_HEX); p3.paragraph_format.left_indent=Pt(8); p3.paragraph_format.space_after=Pt(6)
intro.run(p3,"Plateforme de gestion (back-office d'administration) et portail client B2C/B2B "
             "pour un grossiste alimentaire halal-oriental, région PACA",size=13,color=RGBColor(0xED,0xE7,0xD6))
p4=intro.d.add_paragraph(); _shade(p4._p.get_or_add_pPr(),GREEN_HEX); p4.paragraph_format.left_indent=Pt(8); p4.paragraph_format.space_after=Pt(14)
r=p4.add_run("  Enseigne · Le Bakkal Oriental  "); r.font.name=FONT; r.bold=True; r.font.size=Pt(10)
r.font.color.rgb=GREEN; r.font.highlight_color=WD_COLOR_INDEX.YELLOW
intro.d.add_paragraph()
intro.kv_table([
    ["Commanditaire","Rashed Younes, dirigeant d'Alim Express (PACA)"],
    ["Encadrante pédagogique","Camille SIX"],
    ["Équipe","David · Anas · Yassin"],
    ["Période","3 mars → 22 juin 2026 — soutenance le 1er juillet 2026"],
    ["Composition","Partie 1 : Back-office · Partie 2 : Portail client"],
])
intro.page_break()

intro.h("Accès aux dépôts et aux applications",2)
intro.callout([("Liens d'accès pour le jury. ",True),
    ("Les deux applications et leurs codes sources sont accessibles aux adresses suivantes.",False)])
intro.h("Applications en ligne",3)
intro.link_table(["Application","URL"],[
    ["Portail client (B2C / B2B)","https://lebakkaloriental.fr/"],
    ["Back-office (administration)","https://admin.lebakkaloriental.fr/"],
])
intro.h("Dépôts de code (Git)",3)
intro.link_table(["Dépôt","URL"],[
    ["Portail client — AlimExpressCatalog","https://github.com/mineboxarabic/albakkaloriental_front"],
    ["Back-office — AlimExpressApp","https://github.com/mineboxarabic/AlimExpressApp"],
])
intro.h("Autres livrables",3)
intro.kv_table([
    ["Tableau de pilotage","Trello (lien partagé)"],
    ["Note de cadrage","docs/doc/Note-de-cadrage-Alim-Express-David.pdf"],
    ["Annexes complémentaires","cf. fin du document"],
])
intro.page_break()
intro.save(DOCDIR+r"\_intro.docx")

# ---------- ANNEXES (espaces vides à compléter à la main) ----------
ann=Doc()
ann.h("Annexes",2)
ann.callout([("À compléter manuellement. ",True),
    ("Les espaces ci-dessous sont réservés : insérez-y vos documents (copier-coller, capture ou "
     "lien) directement dans Word.",False)])
annexes=[
    ("Annexe A — Procès-verbal de recettage",
     "[ Espace réservé : insérer ici le PV de recette signé — périmètre testé, anomalies, corrections, statut final. ]"),
    ("Annexe B — Aspect utilisateur & administrateur",
     "[ Espace réservé : insérer ici les guides d'utilisation — parcours utilisateur (portail) et guide administrateur (back-office). ]"),
    ("Annexe C — Suivi de projet (Excel)",
     "[ Espace réservé : insérer ici le suivi de projet (export / capture du fichier Excel : backlog, sprints, planning, avancement). ]"),
    ("Annexe D — Plan de tests",
     "[ Espace réservé : insérer ici le plan de tests détaillé (scénarios, cas de test, résultats attendus/obtenus). ]"),
    ("Annexe E — Aspect technique",
     "[ Espace réservé : insérer ici les éléments techniques complémentaires (schémas d'architecture, modèle de données, contrat d'API, configuration). ]"),
]
for i,(title,hint) in enumerate(annexes):
    ann.h(title,3)
    ann.empty_box(hint,lines=14)
    if i<len(annexes)-1: ann.page_break()
ann.save(DOCDIR+r"\_annexes.docx")

# ---------- FUSION ----------
master=Document(DOCDIR+r"\_intro.docx")
comp=Composer(master)
comp.append(Document(DOCDIR+r"\Dossier-technique-BackOffice-AlimExpress.docx"))
master.add_page_break()
comp.append(Document(DOCDIR+r"\Dossier-technique-Portail-Client-AlimExpress.docx"))
master.add_page_break()
comp.append(Document(DOCDIR+r"\_annexes.docx"))
out=DOCDIR+r"\Dossier-technique-Alim-Express-COMPLET.docx"
comp.save(out)
print("OK ->",out)
